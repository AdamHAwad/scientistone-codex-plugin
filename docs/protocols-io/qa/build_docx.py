#!/usr/bin/env python3
"""Build the review DOCX from the Protocols.io Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "scientist1-beginner-protocol.md"
OUTPUT = ROOT / "scientist1-beginner-protocol.docx"

NAVY = "1D3230"
TEAL = "2E6B62"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GREEN_PALE = "EAF2EE"
SAND = "F6F3EC"
GRAY = "5A6460"
LIGHT_GRAY = "F3F5F4"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(paragraph, fill: str):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def left_border(paragraph, color: str, size="18", space="8"):
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    borders.append(left)


def set_cell_margins(cell, top=180, start=240, bottom=180, end=240):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        margin = tc_mar.find(qn(f"w:{name}"))
        if margin is None:
            margin = OxmlElement(f"w:{name}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def style_callout_cell(cell, fill: str, border: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)

    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "8")
    left.set(qn("w:color"), border)
    borders.append(left)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_spacer(document: Document, points: float):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.line_spacing = Pt(1)
    set_font(paragraph.add_run(" "), size=1, color="FFFFFF")


def add_callout(document: Document, text: str, fill: str, border: str, *, code=False, center=False, width=None):
    add_spacer(document, 2)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER if center else WD_TABLE_ALIGNMENT.LEFT
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    row_properties.append(cannot_split)
    # These single-row tables are self-contained callouts. Marking their only
    # row as a header gives assistive tools an explicit table structure.
    table_header = OxmlElement("w:tblHeader")
    row_properties.append(table_header)
    if width is not None:
        table.autofit = False
        table.columns[0].width = Inches(width)
        table.cell(0, 0).width = Inches(width)
    cell = table.cell(0, 0)
    style_callout_cell(cell, fill, border)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.18
    if code:
        set_font(paragraph.add_run(text), name="Consolas", size=9.5, color=NAVY)
    else:
        add_inline(paragraph, text)
    add_spacer(document, 5)
    return table


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE = re.compile(r"(\*\*.*?\*\*|`.*?`|\[[^]]+\]\([^)]+\))")


def add_inline(paragraph, text: str):
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            set_font(paragraph.add_run(text[position:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), bold=True)
        elif token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), name="Consolas", size=9.5, color=NAVY)
        else:
            link = re.fullmatch(r"\[([^]]+)\]\(([^)]+)\)", token)
            add_hyperlink(paragraph, link.group(1), link.group(2))
        position = match.end()
    if position < len(text):
        set_font(paragraph.add_run(text[position:]))


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=GRAY)


def new_numbering_id(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_properties.extend([tabs, indent])
    level.extend([start, number_format, level_text, justification, paragraph_properties])
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    numbering.append(number)
    return num_id


def new_bullet_numbering_id(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_properties.extend([tabs, indent])
    level.extend([start, number_format, level_text, suffix, justification, paragraph_properties])
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    numbering.append(number)
    return num_id


def set_numbering(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def configure_styles(document: Document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    settings = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in settings.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = rgb(GRAY)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)

    code = styles.add_style("S1 Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(9.5)
    code.font.color.rgb = rgb(NAVY)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.1


def configure_page(document: Document):
    for section in document.sections:
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_after = Pt(0)
        set_font(header.add_run("Scientist1 beginner protocol"), size=9, color=GRAY, bold=True)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.paragraph_format.space_before = Pt(0)
        set_font(footer.add_run("Page "), size=9, color=GRAY)
        add_page_field(footer)


def add_cover(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("BEGINNER PROTOCOL"), size=10, color=TEAL, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("Run a checked research study\nwith Scientist1 in Codex"), size=26, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("From first sign-in to a checked practice study"), size=13, color=GRAY)

    cover = ROOT / "assets" / "00-cover.png"
    pic = document.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = pic.add_run().add_picture(str(cover), width=Inches(6.1))
    shape._inline.docPr.set("descr", "A researcher guides several AI research roles from a question to checked evidence and a paper.")
    pic.paragraph_format.space_after = Pt(14)

    note_table = add_callout(
        document,
        "No coding experience is needed. You will review the plan and any decision that reaches you.",
        GREEN_PALE,
        TEAL,
        center=True,
        width=5.35,
    )
    note = note_table.cell(0, 0).paragraphs[0]
    for run in note.runs:
        set_font(run, size=11, color=NAVY, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(6)
    meta.paragraph_format.space_after = Pt(0)
    set_font(meta.add_run("Version 1.0  |  Checked August 27, 2026"), size=10, color=GRAY)
    document.add_page_break()


def add_image(document: Document, alt: str, relative: str):
    path = ROOT / relative
    with Image.open(path) as image:
        width_px, height_px = image.size
    max_width = 6.2
    max_height = 7.0
    ratio = min(max_width / width_px, max_height / height_px)
    width = max(1.0, width_px * ratio)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True
    if relative.endswith("02-purpose.png"):
        paragraph.paragraph_format.page_break_before = True
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(caption.add_run(alt), size=9, color=GRAY)


def add_paragraph(document: Document, text: str, style=None):
    if text.startswith("Why this"):
        return add_callout(document, text, GREEN_PALE, TEAL)
    paragraph = document.add_paragraph(style=style)
    if text.rstrip().endswith(":"):
        paragraph.paragraph_format.keep_with_next = True
    if text == "Ask:":
        paragraph.paragraph_format.page_break_before = True
    if text in {
        "Keep the page open.",
        "Return to the Codex chat only if it asks you a question.",
        "Do not send the setup form again.",
    }:
        paragraph.paragraph_format.keep_with_next = True
    add_inline(paragraph, text)
    return paragraph


def build():
    document = Document()
    configure_styles(document)
    configure_page(document)
    add_cover(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    first_h1_seen = False
    current_number_id = None
    bullet_number_id = new_bullet_numbering_id(document)

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_callout(document, "\n".join(code_lines), LIGHT_GRAY, TEAL, code=True)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue

        image = re.fullmatch(r"!\[([^]]+)\]\(([^)]+)\)", line)
        if image:
            add_image(document, image.group(1), image.group(2))
            continue

        if line.startswith("# "):
            if not first_h1_seen:
                first_h1_seen = True
                continue
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline(paragraph, line[2:])
            continue
        if line.startswith("## "):
            title = line[3:]
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline(paragraph, title)
            continue
        if line.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline(paragraph, line[4:])
            continue
        if line.startswith("#### "):
            paragraph = document.add_paragraph(style="Heading 3")
            add_inline(paragraph, line[5:])
            continue
        if line.lstrip().startswith("> "):
            add_callout(document, line.lstrip()[2:], SAND, TEAL)
            continue
        if re.match(r"^\s*-\s+", line):
            paragraph = add_paragraph(document, re.sub(r"^\s*-\s+", "", line), style="List Bullet")
            set_numbering(paragraph, bullet_number_id)
            continue
        number_match = re.match(r"^\s*(\d+)\.\s+", line)
        if number_match:
            if current_number_id is None or int(number_match.group(1)) == 1:
                current_number_id = new_numbering_id(document)
            paragraph = add_paragraph(document, re.sub(r"^\s*\d+\.\s+", "", line), style="List Number")
            set_numbering(paragraph, current_number_id)
            continue
        add_paragraph(document, line.strip())

    document.core_properties.title = "Run a checked research study with Scientist1 in Codex"
    document.core_properties.subject = "Beginner Protocols.io guide for Scientist1"
    document.core_properties.keywords = "Scientist1, Codex, research workflow, beginner guide"
    document.core_properties.author = "Scientist1 contributors"
    document.core_properties.last_modified_by = "Scientist1 contributors"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
